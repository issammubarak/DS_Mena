import pandas as pd
from docx import Document

df = pd.read_csv(r'C:\DS4A\DS\Results\DS_MENA_FINAL.csv')
top_30_companies = df['company'].value_counts().head(30).reset_index()
top_30_companies.columns = ['company', 'count']
top_10_companies_per_country = df.groupby('location')['company'].value_counts().groupby(level=0).head(10).reset_index(name='count')
jobs_per_company = df.groupby('company')['job_title'].unique().reset_index()
jobs_per_country = df.groupby('location')['job_title'].unique().reset_index()
jobs_per_country_count = df['location'].value_counts().reset_index()
jobs_per_country_count.columns = ['location', 'count']
total_jobs = len(df)
tools_per_job = df.groupby('job_title')['tools'].unique().reset_index()
tools_per_country = df.groupby('location')['tools'].unique().reset_index()
tools_list = df['tools'].str.split(',').explode()
tools_count = tools_list.value_counts().reset_index()
tools_count.columns = ['tool', 'count']

doc = Document()

def add_dataframe_to_doc(df, doc, title):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    for col_num, column_name in enumerate(df.columns):
        table.cell(0, col_num).text = column_name
    for row_num in range(df.shape[0]):
        for col_num in range(df.shape[1]):
            table.cell(row_num + 1, col_num).text = str(df.iloc[row_num, col_num])

add_dataframe_to_doc(top_30_companies, doc, 'Top 30 companies hiring in all countries')
add_dataframe_to_doc(top_10_companies_per_country, doc, 'Top 10 companies hiring in each country')
add_dataframe_to_doc(jobs_per_company, doc, 'Job titles requested for each company')
add_dataframe_to_doc(jobs_per_country, doc, 'Jobs requested for each country')
add_dataframe_to_doc(jobs_per_country_count, doc, 'Number of jobs requested in each country')

doc.add_heading('Total number of jobs', level=1)
doc.add_paragraph(f'Total jobs: {total_jobs}')
add_dataframe_to_doc(tools_per_job, doc, 'Tools requested for each job')
add_dataframe_to_doc(tools_per_country, doc, 'Tools requested in each country')
add_dataframe_to_doc(tools_count, doc, 'How many times each tool is mentioned in the whole tools column')

doc.save(r'C:\DS4A\DS\Results\job_data_analysis.docx')
print("Data has been exported to job_data_analysis.docx")
